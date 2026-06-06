import os
from docx import Document
from docx.shared import Pt, Inches

def create_system_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('留學機票與學生簽證智慧服務系統 - 系統說明文件', 0)
    title.alignment = 1

    # Overview
    doc.add_heading('一、系統概述', level=1)
    doc.add_paragraph(
        '本系統為一款專為留學代辦顧問與準留學生打造的「一站式雙向服務平台」。系統結合了最新的 AI 技術與即時資料庫，'
        '旨在解決傳統留學申請過程中，資訊繁雜、進度難以追蹤、以及顧問管理負擔過重等痛點。透過分離「學生端」與「顧問端」'
        '雙重視角，達到資訊透明化與自動化管理。'
    )

    # Target Audience & Needs
    doc.add_heading('二、目標使用者與需求', level=1)
    p2 = doc.add_paragraph()
    p2.add_run('1. 學生 (準留學生)：').bold = True
    p2.add_run('需要清晰的倒數計時器、任務清單追蹤（如簽證辦理、機票購買），以及隨時能詢問當地政策的虛擬助理。\n')
    p2.add_run('2. 顧問 (代辦人員)：').bold = True
    p2.add_run('需要一個集中的控制台來管理多位學生的進度，掌握每位學生的風險狀態（正常/預警/緊急），並能快速發送通知與起草催辦信件。')

    # Core Features
    doc.add_heading('三、核心功能介紹', level=1)
    doc.add_paragraph('【雙視角即時同步】：學生端更新資料或修改出航日後，顧問端控制台立即同步，無須等待。', style='List Bullet')
    doc.add_paragraph('【AI 虛擬顧問助理】：搭載雙套 AI 模型。學生端擁有專屬留學地的引導助理；顧問端擁有專屬「政策研究助理」，可依據目前點選的學生動態切換國家政策情境。', style='List Bullet')
    doc.add_paragraph('【智慧催辦信件起草】：系統可根據學生的資料備齊百分比、警示級別及顧問備註，一鍵使用 AI 自動生成客製化的提醒信件。', style='List Bullet')
    doc.add_paragraph('【全球名校探索與地圖】：串接維基百科 API 動態抓取海外學校資訊，並整合 3D 互動地球儀，提供沉浸式的視覺體驗。', style='List Bullet')

    # Tech Stack
    doc.add_heading('四、技術架構', level=1)
    doc.add_paragraph('前端框架：React + TypeScript + Vite', style='List Bullet')
    doc.add_paragraph('樣式與動畫：Tailwind CSS + Framer Motion', style='List Bullet')
    doc.add_paragraph('後端與資料庫：Firebase Authentication + Firestore (即時資料庫)', style='List Bullet')
    doc.add_paragraph('AI 整合：Google Gemini API', style='List Bullet')
    doc.add_paragraph('部屬平台：GitHub Pages', style='List Bullet')

    desktop_path = os.path.join(os.path.expanduser("~"), "OneDrive", "桌面")
    if not os.path.exists(desktop_path):
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        
    save_path = os.path.join(desktop_path, "系統說明文件.docx")
    doc.save(save_path)
    print(f"Saved: {save_path}")

def create_reflection_doc():
    doc = Document()
    
    title = doc.add_heading('AI 協作歷程與學習反思報告', 0)
    title.alignment = 1

    # AI Collaboration
    doc.add_heading('一、AI 協作歷程', level=1)
    doc.add_paragraph(
        '在本次專題開發中，我大量使用了 AI (如 Claude / Gemini) 進行協作，主要應用在三個階段：\n'
        '1. 系統架構與介面設計：初期透過與 AI 討論，確立了「雙視角(學生/顧問)」的架構，並讓 AI 協助生成 Tailwind CSS 的版面配置與 Framer Motion 的微動畫，提升視覺豐富度。\n'
        '2. 功能實作與 API 串接：在串接維基百科 API 與 Firebase 即時資料庫時，AI 提供了清晰的非同步處理範例，並幫忙規劃了 React Context 與 State 的管理方式。\n'
        '3. 除錯與最佳化：開發後期遇到許多「狀態不同步」與「頁面空白(如 GitHub Pages 路由問題)」的 Bug，AI 能夠根據錯誤代碼或截圖，迅速給出 Vite base path 修改與 GitHub Actions 腳本建置的解決方案。'
    )

    # Challenges
    doc.add_heading('二、開發挑戰與解決方式', level=1)
    doc.add_paragraph(
        '最大的挑戰在於「資料狀態的同步」與「AI Context 的切換」。一開始，當顧問在後台點擊不同學生時，AI 助手的對話情境並不會跟著切換，甚至會發生資料互相覆蓋（管理員資料蓋過學生資料）的嚴重問題。\n'
        '解決方式：透過與 AI 進行深度的程式碼除錯，我們重新定義了 initialActiveStudent 的傳遞機制，確保在顧問模式下，系統會優先讀取「被選取學生」的國家與大學欄位，而不是全域變數，最終成功讓後台的 AI 助手達到動態切換情境的效果。'
    )

    # Learning
    doc.add_heading('三、學習收穫', level=1)
    doc.add_paragraph(
        '1. 掌握現代 Web 開發技術棧：從無到有建置了 React + Vite 的專案，並深入了解了 Firebase 的 NoSQL 資料庫運作原理與即時監聽機制。\n'
        '2. Prompt Engineering (提示工程)：學會如何精確地對 AI 下達指令。當程式碼超過上千行時，發現不能只是丟問題給 AI，而是要清楚描述「目前架構」、「預期結果」以及「錯誤訊息」，AI 才能給出對的修改建議。\n'
        '3. 專案部屬經驗：熟悉了從本地端打包 (npm run build) 到透過 GitHub Actions 自動化 CI/CD 部屬至 GitHub Pages 的完整流程。'
    )

    # Future Outlook
    doc.add_heading('四、未來展望與個人反思', level=1)
    doc.add_paragraph(
        '這次專題讓我深刻體會到，未來的軟體開發已經不再是「純手工敲程式碼」，而是「與 AI 結對編程 (Pair Programming)」。開發者的核心價值轉變為「系統設計思維」與「需求定義能力」。\n'
        '未來若有機會，我希望能將此系統串接真實的機票 API (如 Skyscanner/Amadeus) 以取得即時票價，並整合 LINE Messaging API，讓系統能夠真正透過通訊軟體發送自動化推播通知，使專案更加完整且具商業價值。'
    )

    desktop_path = os.path.join(os.path.expanduser("~"), "OneDrive", "桌面")
    if not os.path.exists(desktop_path):
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        
    save_path = os.path.join(desktop_path, "AI協作歷程與學習反思報告.docx")
    doc.save(save_path)
    print(f"Saved: {save_path}")

if __name__ == "__main__":
    create_system_doc()
    create_reflection_doc()
